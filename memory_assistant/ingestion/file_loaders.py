"""
文件加载器 - 支持多种文档格式
"""
import os
import hashlib
from abc import ABC, abstractmethod
from typing import Optional, Tuple
from datetime import datetime

from .models import DocumentMetadata, DocumentType


class BaseLoader(ABC):
    """加载器基类"""

    @abstractmethod
    def load(self, file_path: str) -> Tuple[str, dict]:
        """返回 (文本内容, 提取的元数据)"""
        pass

    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """检查是否支持该文件"""
        pass


class PDFLoader(BaseLoader):
    """PDF加载器（仅支持文本型PDF）"""

    def __init__(self):
        try:
            import fitz  # PyMuPDF
            self.fitz = fitz
        except ImportError:
            raise ImportError("请安装 PyMuPDF: pip install PyMuPDF")

    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')

    def load(self, file_path: str) -> Tuple[str, dict]:
        doc = self.fitz.open(file_path)

        # 提取文本
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())

        # 提取元数据
        metadata = {
            'page_count': len(doc),
            'author': doc.metadata.get('author'),
            'title': doc.metadata.get('title'),
            'subject': doc.metadata.get('subject'),
        }

        doc.close()
        return "\n\n".join(text_parts), metadata


class DocxLoader(BaseLoader):
    """Word文档加载器"""

    def __init__(self):
        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith('.docx')

    def load(self, file_path: str) -> Tuple[str, dict]:
        doc = self.Document(file_path)

        # 提取段落
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 提取表格（会议记录可能包含表格）
        tables_text = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_rows.append(" | ".join(row_text))
            tables_text.append("\n".join(table_rows))

        all_text = "\n\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n【表格内容】\n" + "\n\n".join(tables_text)

        # 提取元数据
        metadata = {
            'author': doc.core_properties.author,
            'title': doc.core_properties.title,
            'paragraph_count': len(paragraphs),
        }

        return all_text, metadata


class TextLoader(BaseLoader):
    """纯文本/Markdown加载器"""

    def supports(self, file_path: str) -> bool:
        ext = file_path.lower()
        return ext.endswith('.txt') or ext.endswith('.md') or ext.endswith('.markdown')

    def load(self, file_path: str) -> Tuple[str, dict]:
        # 自动检测编码
        import chardet
        with open(file_path, 'rb') as f:
            raw = f.read()
            encoding = chardet.detect(raw)['encoding'] or 'utf-8'

        content = raw.decode(encoding, errors='ignore')

        # 简单统计
        metadata = {
            'line_count': content.count('\n') + 1,
            'char_count': len(content),
        }

        return content, metadata


class FileLoaderFactory:
    """加载器工厂"""

    _loaders = [
        PDFLoader(),
        DocxLoader(),
        TextLoader(),
    ]

    @classmethod
    def get_loader(cls, file_path: str) -> Optional[BaseLoader]:
        for loader in cls._loaders:
            if loader.supports(file_path):
                return loader
        return None

    @classmethod
    def load_file(cls, file_path: str) -> Tuple[str, dict, DocumentType]:
        loader = cls.get_loader(file_path)
        if not loader:
            raise ValueError(f"不支持的文件格式: {file_path}")

        content, metadata = loader.load(file_path)

        # 确定文档类型
        ext = os.path.splitext(file_path)[1].lower()
        type_map = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN,
        }
        doc_type = type_map.get(ext, DocumentType.TXT)

        return content, metadata, doc_type
